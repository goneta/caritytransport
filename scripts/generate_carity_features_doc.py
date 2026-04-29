from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
DOCX_PATH = OUT_DIR / "Carity Application Features Overview.docx"


ACCENT = RGBColor(13, 59, 63)
ACCENT_SOFT = RGBColor(32, 150, 158)
TEXT = RGBColor(33, 37, 41)
MUTED = RGBColor(102, 112, 122)
LIGHT = RGBColor(242, 245, 247)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    if color:
        run.font.color.rgb = color


def style_run(run, *, size=10.5, bold=False, color: RGBColor | None = None) -> None:
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_body_paragraph(doc: Document, text: str, *, spacing_after=8) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = 1.18
    run = p.add_run(text)
    style_run(run, size=10.5, color=TEXT)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    style_run(run, size=10.2, color=TEXT)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    style_run(run, size=16 if level == 1 else 12.5, bold=True, color=ACCENT if level == 1 else TEXT)


def add_feature_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(2.0), Inches(2.05), Inches(2.65)]
    headers = ["Area", "What the User Can Do", "Business Value"]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width

    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(hdr[i], "0D3B3F")
        set_cell_text(hdr[i], header, bold=True, color=RGBColor(255, 255, 255), size=10)

    for area, what, value in rows:
        row = table.add_row().cells
        set_cell_shading(row[0], "F2F5F7")
        set_cell_text(row[0], area, bold=True, color=ACCENT, size=10)
        set_cell_text(row[1], what, color=TEXT, size=10)
        set_cell_text(row[2], value, color=TEXT, size=10)


def add_note_box(doc: Document, title: str, body: str, fill: str = "EEF6F6") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.2)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(title)
    style_run(r1, size=10.5, bold=True, color=ACCENT)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    r2 = p2.add_run(body)
    style_run(r2, size=10, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_doc() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cover.paragraph_format.space_after = Pt(8)
    r = cover.add_run("Carity")
    style_run(r, size=24, bold=True, color=ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run("Application Features Overview")
    style_run(r, size=19, bold=True, color=TEXT)

    strap = doc.add_paragraph()
    strap.paragraph_format.space_after = Pt(12)
    strap.paragraph_format.line_spacing = 1.2
    r = strap.add_run(
        "Client-facing summary of the school transport workflows, dashboards, and operational controls currently represented in the Carity codebase."
    )
    style_run(r, size=11, color=MUTED)

    add_note_box(
        doc,
        "Executive Summary",
        "Carity is a role-based school transport platform designed for administrators, schedulers, drivers, parents, and pupils. The application covers route creation, seat booking, pupil management, QR-based boarding checks, notifications, reporting, compliance document storage, customer support workflows, and recruitment administration.",
        fill="F4F8F8",
    )

    summary_table = doc.add_table(rows=2, cols=2)
    summary_table.autofit = False
    summary_table.columns[0].width = Inches(2.1)
    summary_table.columns[1].width = Inches(4.1)
    items = [
        ("Primary users", "Admin teams, drivers, parents, and pupils"),
        ("Core purpose", "Manage school transport safely from registration through daily operations"),
    ]
    for row_idx, (left, right) in enumerate(items):
        row = summary_table.rows[row_idx].cells
        set_cell_shading(row[0], "F2F5F7")
        set_cell_text(row[0], left, bold=True, color=ACCENT, size=10)
        set_cell_text(row[1], right, color=TEXT, size=10)

    doc.add_page_break()

    add_heading(doc, "Platform At a Glance", 1)
    add_body_paragraph(
        doc,
        "The platform is organised around separate dashboards for each operational role. This makes the application useful both as a parent-facing service and as an internal transport operations system."
    )
    add_feature_table(
        doc,
        [
            ("Role-based access", "Separate experiences for admin, driver, parent, and pupil users.", "Keeps each user focused on the tasks relevant to them."),
            ("Identity and access", "Login, registration, user profiles, status tracking, and QR identity cards.", "Supports secure access and traceable user activity."),
            ("Operational data model", "Central records for schools, pupils, parents, vehicles, drivers, bookings, payments, documents, and support cases.", "Creates one shared source of truth across the transport workflow."),
        ],
    )

    add_heading(doc, "Administrative Features", 1)
    add_feature_table(
        doc,
        [
            ("Dashboard and alerts", "View totals for pupils, routes, vehicles, drivers, parents, recent activity, and pending actions such as expiring licences or insurance.", "Gives leadership a live operating snapshot."),
            ("People management", "Maintain parent, pupil, and employee records, including contact details, status, transport eligibility, and profile information.", "Improves onboarding and day-to-day record keeping."),
            ("Fleet management", "Manage transport companies, drivers, vehicles, MOT records, insurance, and licence classes.", "Supports safety, compliance, and contractor oversight."),
            ("Route scheduling", "Create, edit, clone, and monitor routes with service type, recurrence, school links, postcodes, pricing, driver assignment, and vehicle capacity.", "Makes route planning operationally practical and scalable."),
            ("Capacity and utilisation", "See assigned seats against vehicle capacity and highlight full or high-utilisation routes.", "Helps prevent overbooking and improves planning decisions."),
            ("Bookings oversight", "Review confirmed transport bookings and booking activity through admin APIs and dashboard metrics.", "Improves financial and service visibility."),
            ("Resolution centre", "Handle complaints and refund requests, update ticket status, add parent replies, and process refund outcomes.", "Provides a structured support workflow."),
            ("Notifications", "Trigger and track in-app, email, and SMS-style notification preferences and delivery records.", "Improves communication with families and staff."),
            ("Analytics", "Charts for revenue, booking status, route utilisation, fleet mix, and driver status.", "Turns operational data into management insight."),
            ("Document vault", "Upload and track DBS, insurance, MOT, and licence documents with expiry dates linked to drivers or vehicles.", "Strengthens compliance management."),
            ("Recruitment module", "Review job applications, schedule interviews, hire candidates, create employee credentials, and keep internal notes.", "Extends the platform into staff recruitment and onboarding."),
            ("Audit trail", "Store audit log records of changes across the platform.", "Supports accountability and administrative traceability."),
        ],
    )

    add_heading(doc, "Parent Experience", 1)
    add_feature_table(
        doc,
        [
            ("Parent dashboard", "See registered children, active routes, unread alerts, upcoming absences, and a personal QR identity card.", "Provides a simple command centre for families."),
            ("Child management", "Register children, store school details, special requirements, emergency contacts, and transport status.", "Helps the service capture safeguarding and operational information."),
            ("Transport search", "Search for available routes by postcode, school, trip type, date, and preferred time.", "Makes discovery easy for families."),
            ("Seat selection", "Choose a vehicle, view seat availability, and assign a seat to a specific child before booking.", "Creates a transparent self-service booking flow."),
            ("Basket and checkout", "Build a basket of journeys, review totals, and move through a protected payment-confirmation flow with seat conflict checks.", "Reduces booking errors and improves purchase confidence."),
            ("Bookings and history", "Access booked trips, historical transport records, and schedule information.", "Supports day-to-day travel management."),
            ("Live tracking screen", "View route progress, vehicle details, and timeline-style journey status updates for each child.", "Improves reassurance and visibility for parents."),
            ("Notifications", "Receive route, booking, and journey updates in a dedicated notification area.", "Keeps families informed without calling support."),
            ("Resolution centre", "Raise refund requests, complaints, or other support issues, and continue the conversation with support staff.", "Creates a clear customer service path."),
        ],
    )

    add_heading(doc, "Driver Experience", 1)
    add_feature_table(
        doc,
        [
            ("Driver dashboard", "See assigned routes, total pupils, schedule summaries, and a personal QR identity card.", "Gives drivers a clear start-of-day view."),
            ("Schedule and route views", "Access route lists, route timing, vehicle details, and assigned passenger counts.", "Supports route readiness and execution."),
            ("Manifest and attendance", "Review assigned pupils and mark boarded or absent status from the route view.", "Helps drivers manage day-of-service operations."),
            ("QR boarding scanner", "Scan or paste pupil QR data, validate whether the pupil is booked, and display emergency or parent contact details.", "Improves boarding control and safeguarding confidence."),
            ("Trip logging", "Log boarding and route events to trip records when QR scans occur.", "Creates a record of operational activity."),
            ("Unavailability reporting", "Driver data model and APIs support absence reporting and substitute workflows.", "Reduces service disruption when staff availability changes."),
        ],
    )

    add_heading(doc, "Pupil and Support Features", 1)
    add_feature_table(
        doc,
        [
            ("Pupil portal", "Provide pupil-facing dashboard and trip views linked to the child account.", "Extends visibility beyond the parent account."),
            ("AI assistant", "Offer contextual chat support for parents, drivers, and administrators, using role-aware prompts and stored conversation history.", "Improves support responsiveness and self-service guidance."),
            ("Knowledge support", "Answer questions about routes, bookings, company information, and careers.", "Helps reduce repetitive enquiries."),
            ("Notifications and messaging", "Store outbound notifications and support ticket replies in the platform database.", "Keeps communication traceable."),
        ],
    )

    add_heading(doc, "Compliance, Safety, and Control", 1)
    add_bullet(doc, "Driver, vehicle, insurance, DBS, and licence records are all represented in the core data model.")
    add_bullet(doc, "QR identity and boarding flows help confirm whether a pupil is expected on a route.")
    add_bullet(doc, "Trip logs, audit logs, and support tickets provide traceability across operations.")
    add_bullet(doc, "Notification preferences exist for SMS, email, and push-style communications.")
    add_bullet(doc, "Holiday periods, absences, and driver unavailability support real-world transport exceptions.")

    add_heading(doc, "Implementation Notes", 1)
    add_note_box(
        doc,
        "Important Context for Client Conversations",
        "The codebase includes the full product structure for payments, live tracking, and AI support, but some flows still depend on configuration or further production wiring. For example, the tracking screen currently shows a map placeholder until a maps key is configured, and the checkout route presently simulates a successful payment rather than creating a live Stripe checkout session. These do not change the product direction, but they are worth presenting as implementation-stage items rather than fully deployed production integrations.",
        fill="FFF7E8",
    )

    add_heading(doc, "Pricing and External Services", 1)
    add_body_paragraph(
        doc,
        "The proposed price for the Carity web application is £5,000. This covers the main platform scope represented in the application, including administration, parent and driver dashboards, route scheduling, seat booking, QR boarding, notifications, reporting, compliance document handling, support workflows, and recruitment features."
    )
    add_feature_table(
        doc,
        [
            ("Platform price", "Full Carity web application.", "£5,000 total."),
            ("AI chatbot", "Uses an external AI model provider for conversational support.", "Client subscription required with providers such as Anthropic or OpenAI."),
            ("Text messaging", "SMS notifications require an external messaging provider.", "Client subscription required with Twilio."),
            ("Online payments", "Stripe is required for live card payment processing.", "Client must create and maintain a Stripe account."),
            ("Domain and hosting", "Live deployment infrastructure and domain registration.", "External charge of £45, not included in the £5,000 platform price."),
        ],
    )
    add_note_box(
        doc,
        "Client Setup Requirements",
        "To operate the live service, the client will need to create their own Stripe account for online payments, an AI account with a provider such as Anthropic or OpenAI for chatbot usage, and a Twilio account for SMS messaging. These third-party subscriptions and usage charges are separate from the £5,000 web application price.",
        fill="EEF6F6",
    )

    add_heading(doc, "Conclusion", 1)
    add_body_paragraph(
        doc,
        "Based on the current codebase, Carity is positioned as a comprehensive school transport management application rather than a simple booking portal. It combines family self-service, operational control, driver tooling, compliance records, and support workflows in one platform."
    )
    add_body_paragraph(
        doc,
        "This document was prepared from the implemented routes, pages, APIs, and data model in the repository on 27 April 2026."
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
